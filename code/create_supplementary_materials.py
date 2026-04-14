from pathlib import Path
import json
import re
import shutil
import textwrap

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
SRC_BASE = ROOT / "output"
SRC_FIG = SRC_BASE / "figures"
SRC_TBL = SRC_BASE / "tables"

OUT_BASE = ROOT / "supplementary_materials"
OUT_FIG = OUT_BASE / "figures"
OUT_TBL = OUT_BASE / "tables"
OUT_DOCX = OUT_BASE / "BMC_Ophthalmology_Supplementary_Material.docx"


def ensure_dirs():
    OUT_FIG.mkdir(parents=True, exist_ok=True)
    OUT_TBL.mkdir(parents=True, exist_ok=True)


def copy_figure(src_name: str, dst_name: str) -> Path:
    src = SRC_FIG / src_name
    dst = OUT_FIG / dst_name
    shutil.copy2(src, dst)
    return dst


def normalize_for_table(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out.columns = [c.replace("_", " ").replace("BAD D", "BAD-D") for c in out.columns]
    replace_map = {
        "K Max (Front)": "Kmax",
        "Strong separation": "Strong separation",
        "Strong separation": "Strong separation",
        "Artificial split": "Artificial split",
        "Artificial split": "Artificial split",
        "none": "None",
    }
    for col in out.columns:
        if out[col].dtype == object:
            out[col] = out[col].astype(str)
            if col.lower() in {"cluster sizes", "sizes"}:
                out[col] = out[col].str.replace(r"np\.int64\((\d+)\)", r"\1", regex=True)
                out[col] = out[col].str.replace("[", "", regex=False).str.replace("]", "", regex=False)
                out[col] = out[col].apply(lambda x: textwrap.fill(x, width=16))
            for old, new in replace_map.items():
                out[col] = out[col].str.replace(old, new, regex=False)
            if col.lower() in {"nonsig vars", "iqr"}:
                out[col] = out[col].apply(lambda x: textwrap.fill(x, width=26))
    return out


def render_dataframe_image(
    df: pd.DataFrame,
    out_path: Path,
    title: str,
    subtitle: str | None = None,
    font_size: int = 9,
    row_height: float = 0.36,
):
    df = normalize_for_table(df)
    n_rows, n_cols = df.shape
    fig_w = max(12, min(20, 1.55 * n_cols + 1.2))
    fig_h = max(2.2, 1.25 + row_height * (n_rows + 1))
    fig, ax = plt.subplots(figsize=(fig_w, fig_h), dpi=300)
    ax.axis("off")

    title_y = 0.98
    fig.text(0.5, title_y, title, ha="center", va="top", fontsize=14, fontweight="bold")
    if subtitle:
        fig.text(0.5, title_y - 0.04, subtitle, ha="center", va="top", fontsize=9)

    table = ax.table(
        cellText=df.values,
        colLabels=df.columns,
        loc="center",
        cellLoc="center",
        colLoc="center",
        bbox=[0.0, 0.0, 1.0, 0.90 if subtitle else 0.93],
    )
    table.auto_set_font_size(False)
    table.set_fontsize(font_size)
    table.scale(1, 1.15)

    for (row, col), cell in table.get_celld().items():
        cell.set_edgecolor("#7a7a7a")
        cell.set_linewidth(0.6)
        if row == 0:
            cell.set_facecolor("#dbe9f6")
            cell.set_text_props(weight="bold")
        else:
            cell.set_facecolor("#ffffff" if row % 2 else "#f7f7f7")

    plt.savefig(out_path, bbox_inches="tight", pad_inches=0.15)
    plt.close(fig)


def render_chunked_table(
    df: pd.DataFrame,
    base_name: str,
    title: str,
    subtitle: str | None,
    rows_per_chunk: int,
    font_size: int = 8,
):
    paths = []
    df = df.reset_index(drop=True)
    total = len(df)
    if total <= rows_per_chunk:
        out_path = OUT_TBL / f"{base_name}.png"
        render_dataframe_image(df, out_path, title, subtitle, font_size=font_size)
        return [out_path]

    for idx, start in enumerate(range(0, total, rows_per_chunk), start=1):
        chunk = df.iloc[start:start + rows_per_chunk].copy()
        chunk_title = f"{title} (continued {idx})" if idx > 1 else title
        out_path = OUT_TBL / f"{base_name}_part{idx}.png"
        render_dataframe_image(chunk, out_path, chunk_title, subtitle if idx == 1 else None, font_size=font_size)
        paths.append(out_path)
    return paths


def compute_table_s1() -> pd.DataFrame:
    df = pd.read_csv(SRC_TBL / "k_comparison_extended.csv")
    metric_cols = ["cosine_silhouette", "bootstrap_ARI_mean", "severity_LR_acc", "max_centroid_sim"]
    normalized = {}
    for col in metric_cols:
        vals = df[col].astype(float).to_numpy()
        vmin, vmax = vals.min(), vals.max()
        scale = (vals - vmin) / (vmax - vmin + 1e-10)
        normalized[col] = 1 - scale if col in {"severity_LR_acc", "max_centroid_sim"} else scale

    composite = (
        0.25 * normalized["cosine_silhouette"]
        + 0.30 * normalized["bootstrap_ARI_mean"]
        + 0.25 * normalized["severity_LR_acc"]
        + 0.20 * normalized["max_centroid_sim"]
    )
    df["composite_score"] = composite.round(4)
    df["selected"] = np.where(df["K"] == 2, "Yes", "")

    out = df[
        [
            "K",
            "sizes",
            "min_cluster_pct",
            "cosine_silhouette",
            "euclidean_silhouette",
            "bootstrap_ARI_mean",
            "bootstrap_ARI_sd",
            "severity_LR_acc",
            "max_centroid_sim",
            "composite_score",
            "selected",
        ]
    ].copy()
    out.columns = [
        "K",
        "Cluster sizes",
        "Min cluster %",
        "Cosine silhouette",
        "Euclidean silhouette",
        "Bootstrap ARI",
        "Bootstrap ARI SD",
        "Severity LR acc",
        "Max centroid sim",
        "Composite score",
        "Selected",
    ]
    return out


def compute_table_s6() -> pd.DataFrame:
    with open(SRC_TBL / "centroid_concordance.json", "r", encoding="utf-8") as f:
        payload = json.load(f)

    matrix = payload["centroid_similarity_matrix"]
    matched = {disc: repl for disc, repl in payload["matched_pairs_disc_to_repl"]}
    matched_scores = payload["matched_similarities"]

    rows = []
    for disc_idx, row in enumerate(matrix):
        repl_idx = matched[disc_idx]
        rows.append(
            {
                "Discovery centroid": f"P{disc_idx + 1}",
                "Replication P1": round(row[0], 4),
                "Replication P2": round(row[1], 4),
                "Replication P3": round(row[2], 4),
                "Replication P4": round(row[3], 4),
                "Best match": f"P{repl_idx + 1}",
                "Matched cosine similarity": round(row[repl_idx], 4),
            }
        )

    rows.append(
        {
            "Discovery centroid": "Mean matched similarity",
            "Replication P1": "",
            "Replication P2": "",
            "Replication P3": "",
            "Replication P4": "",
            "Best match": "",
            "Matched cosine similarity": round(payload["mean_concordance"], 4),
        }
    )
    rows.append(
        {
            "Discovery centroid": "ARI (projected vs refit)",
            "Replication P1": "",
            "Replication P2": "",
            "Replication P3": "",
            "Replication P4": "",
            "Best match": "",
            "Matched cosine similarity": round(payload["ARI_projected_vs_independent"], 4),
        }
    )
    return pd.DataFrame(rows)


def compute_table_s7() -> pd.DataFrame:
    df = pd.read_csv(SRC_TBL / "weight_sensitivity.csv")
    keep = [
        "weight_scheme",
        "composite_K2",
        "composite_K3",
        "composite_K4",
        "composite_K5",
        "composite_K6",
        "optimal_K",
    ]
    out = df[keep].copy()
    out.columns = [
        "Weight scheme",
        "K = 2",
        "K = 3",
        "K = 4",
        "K = 5",
        "K = 6",
        "Optimal K",
    ]
    return out


def compute_table_s8() -> pd.DataFrame:
    df = pd.read_csv(SRC_TBL / "silhouette_comparison.csv")
    out = df.copy()
    out.columns = ["K", "Cosine silhouette", "Euclidean silhouette", "Difference"]
    out["K"] = out["K"].astype(int)
    return out


def create_figure_s2(out_path: Path):
    df = pd.read_csv(SRC_TBL / "k_comparison_extended.csv")
    metric_cols = {
        "Cosine silhouette": df["cosine_silhouette"].astype(float),
        "Bootstrap ARI": df["bootstrap_ARI_mean"].astype(float),
        "Severity LR accuracy": df["severity_LR_acc"].astype(float),
        "Max centroid similarity": df["max_centroid_sim"].astype(float),
    }

    norm_sil = (metric_cols["Cosine silhouette"] - metric_cols["Cosine silhouette"].min()) / (
        metric_cols["Cosine silhouette"].max() - metric_cols["Cosine silhouette"].min() + 1e-10
    )
    norm_stab = (metric_cols["Bootstrap ARI"] - metric_cols["Bootstrap ARI"].min()) / (
        metric_cols["Bootstrap ARI"].max() - metric_cols["Bootstrap ARI"].min() + 1e-10
    )
    norm_sev = 1 - (metric_cols["Severity LR accuracy"] - metric_cols["Severity LR accuracy"].min()) / (
        metric_cols["Severity LR accuracy"].max() - metric_cols["Severity LR accuracy"].min() + 1e-10
    )
    norm_sim = 1 - (metric_cols["Max centroid similarity"] - metric_cols["Max centroid similarity"].min()) / (
        metric_cols["Max centroid similarity"].max() - metric_cols["Max centroid similarity"].min() + 1e-10
    )
    composite = 0.25 * norm_sil + 0.30 * norm_stab + 0.25 * norm_sev + 0.20 * norm_sim

    fig, axes = plt.subplots(2, 3, figsize=(14, 8), dpi=300)
    axes = axes.ravel()
    ks = df["K"].astype(int).tolist()
    panels = [
        ("Cosine silhouette", metric_cols["Cosine silhouette"], "#4C78A8"),
        ("Bootstrap ARI", metric_cols["Bootstrap ARI"], "#59A14F"),
        ("Severity LR accuracy", metric_cols["Severity LR accuracy"], "#F28E2B"),
        ("Max centroid similarity", metric_cols["Max centroid similarity"], "#9C3DA5"),
        ("Composite score", composite, "#E15759"),
    ]

    for ax, (title, values, color) in zip(axes, panels):
        ax.plot(ks, values, marker="o", linewidth=2.2, color=color)
        ax.set_title(title, fontsize=11, fontweight="bold")
        ax.set_xticks(ks)
        ax.grid(alpha=0.25, linestyle="--")
        for x, y in zip(ks, values):
            ax.text(x, y, f"{y:.3f}", fontsize=8, ha="center", va="bottom")
        if title == "Composite score":
            best_idx = int(np.argmax(values))
            ax.axvline(ks[best_idx], color="red", linestyle="--", linewidth=1)
            ax.text(ks[best_idx], max(values), "Selected K=2", color="red", fontsize=8, ha="left", va="top")

    axes[5].axis("off")
    fig.suptitle("Supplementary Figure S2. Cluster-number selection metrics across K = 2 to 6", fontsize=14, fontweight="bold")
    fig.tight_layout(rect=[0, 0, 1, 0.95])
    fig.savefig(out_path, bbox_inches="tight", pad_inches=0.15)
    plt.close(fig)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.add_run(text)
    return p


def add_image(doc: Document, path: Path, width: float, caption: str):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(str(path), width=Inches(width))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_cap.add_run(caption)
    r.bold = False


def build_docx(fig_paths: dict[str, Path], table_paths: dict[str, list[Path]]):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Supplementary Material")
    r.bold = True
    r.font.size = doc.styles["Title"].font.size

    add_heading(doc, "Supplementary Methods", 1)
    doc.add_paragraph(
        "Additional sensitivity analyses were performed using the same filtered cohort, dual-surface "
        "normalization framework, and discovery-replication split described in the main manuscript. "
        "The Supplementary Material reports the prespecified spherical k-means cluster-number metrics "
        "across K = 2 to 6 (Table S1; Figure S2), the full-feature and HOA-only K = 4 over-segmentation "
        "probe with post-hoc clinical summaries and pairwise distinguishability analyses (Figures S3-S4; "
        "Tables S2-S5), discovery-replication centroid concordance after centroid matching (Table S6), "
        "composite-score robustness to alternative weighting schemes (Table S7), and "
        "cosine-versus-Euclidean silhouette behavior across K = 2 to 6 (Table S8). The HOA-only feature "
        "set was used specifically to test whether the K = 4 over-segmentation pattern persisted in an "
        "independent subset of aberration features."
    )

    add_heading(doc, "Supplementary Figures", 1)
    add_image(
        doc,
        fig_paths["S1"],
        6.8,
        "Supplementary Figure S1. Cumulative explained variance of PCA fitted on the discovery set "
        "for the full dual-surface feature set (left) and the HOA-only feature set (right). "
        "Vertical dashed lines indicate the retained 20-component threshold; horizontal dashed lines "
        "indicate the 90% and 95% reference levels.",
    )
    add_image(
        doc,
        fig_paths["S2"],
        6.8,
        "Supplementary Figure S2. Cluster-number selection metrics across K = 2 to 6 in the discovery set. "
        "Panels show cosine silhouette, bootstrap stability, severity predictability from total RMS alone, "
        "maximum pairwise centroid similarity, and the composite score used for cluster-number selection.",
    )
    add_image(
        doc,
        fig_paths["S3"],
        6.3,
        "Supplementary Figure S3. Full-feature K = 4 over-segmentation atlas. The four-group probe was used "
        "to test whether forcing additional resolution produced clinically distinct morphologies; the associated "
        "pairwise post-hoc comparisons are summarized in Supplementary Table S3.",
    )
    add_image(
        doc,
        fig_paths["S4"],
        6.3,
        "Supplementary Figure S4. HOA-only K = 4 over-segmentation atlas. The same probe was repeated using "
        "orders 3 to 6 only, providing an independent sensitivity analysis of the four-group solution.",
    )

    add_heading(doc, "Supplementary Tables", 1)
    captions = {
        "S1": "Supplementary Table S1. Cluster-number comparison across K = 2 to 6 for the primary spherical k-means analysis in the discovery set.",
        "S2": "Supplementary Table S2. Post-hoc clinical summary of the full-feature K = 4 over-segmentation probe.",
        "S3": "Supplementary Table S3. Pairwise distinguishability analysis for the full-feature K = 4 solution.",
        "S4": "Supplementary Table S4. Post-hoc clinical summary of the HOA-only K = 4 over-segmentation probe.",
        "S5": "Supplementary Table S5. Pairwise distinguishability analysis for the HOA-only K = 4 solution.",
        "S6": "Supplementary Table S6. Discovery-replication centroid concordance after centroid matching for the K = 4 over-segmentation probe.",
        "S7": "Supplementary Table S7. Composite-score sensitivity to alternative weighting schemes across K = 2 to 6.",
        "S8": "Supplementary Table S8. Cosine-versus-Euclidean silhouette comparison across K = 2 to 6.",
    }
    for label in ["S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8"]:
        first = True
        for idx, img_path in enumerate(table_paths[label], start=1):
            caption = captions[label] if first else f"{captions[label]} (continued)"
            add_image(doc, img_path, 6.8, caption)
            first = False

    doc.save(str(OUT_DOCX))


def main():
    ensure_dirs()

    figure_paths = {
        "S1": copy_figure("fig02_pca_scree.png", "Supplementary_Figure_S1.png"),
        "S3": copy_figure("fig04_full_K4.png", "Supplementary_Figure_S3.png"),
        "S4": copy_figure("fig05_hoa_K4.png", "Supplementary_Figure_S4.png"),
    }
    figure_paths["S2"] = OUT_FIG / "Supplementary_Figure_S2.png"
    create_figure_s2(figure_paths["S2"])

    table_paths = {}

    table_s1 = compute_table_s1()
    table_paths["S1"] = render_chunked_table(
        table_s1,
        "Supplementary_Table_S1",
        "Supplementary Table S1",
        "Primary spherical k-means cluster-number comparison across K = 2 to 6 in the discovery set",
        rows_per_chunk=12,
        font_size=8,
    )

    full_summary = pd.read_csv(SRC_TBL / "clinical_summary_full.csv")
    table_paths["S2"] = render_chunked_table(
        full_summary,
        "Supplementary_Table_S2",
        "Supplementary Table S2",
        "Post-hoc clinical summary of the full-feature K = 4 over-segmentation probe",
        rows_per_chunk=12,
        font_size=8,
    )

    pairwise_full = pd.read_csv(SRC_TBL / "pairwise_full.csv")
    table_paths["S3"] = render_chunked_table(
        pairwise_full,
        "Supplementary_Table_S3",
        "Supplementary Table S3",
        "Pairwise distinguishability analysis for the full-feature K = 4 solution",
        rows_per_chunk=12,
        font_size=8,
    )

    hoa_summary = pd.read_csv(SRC_TBL / "clinical_summary_hoa.csv")
    table_paths["S4"] = render_chunked_table(
        hoa_summary,
        "Supplementary_Table_S4",
        "Supplementary Table S4",
        "Post-hoc clinical summary of the HOA-only K = 4 over-segmentation probe",
        rows_per_chunk=12,
        font_size=8,
    )

    pairwise_hoa = pd.read_csv(SRC_TBL / "pairwise_hoa.csv")
    table_paths["S5"] = render_chunked_table(
        pairwise_hoa,
        "Supplementary_Table_S5",
        "Supplementary Table S5",
        "Pairwise distinguishability analysis for the HOA-only K = 4 solution",
        rows_per_chunk=12,
        font_size=8,
    )

    table_s6 = compute_table_s6()
    table_paths["S6"] = render_chunked_table(
        table_s6,
        "Supplementary_Table_S6",
        "Supplementary Table S6",
        "Discovery-replication centroid concordance after centroid matching for the K = 4 over-segmentation probe",
        rows_per_chunk=12,
        font_size=8,
    )

    table_s7 = compute_table_s7()
    table_paths["S7"] = render_chunked_table(
        table_s7,
        "Supplementary_Table_S7",
        "Supplementary Table S7",
        "Composite-score sensitivity to alternative weighting schemes across K = 2 to 6",
        rows_per_chunk=12,
        font_size=8,
    )

    table_s8 = compute_table_s8()
    table_paths["S8"] = render_chunked_table(
        table_s8,
        "Supplementary_Table_S8",
        "Supplementary Table S8",
        "Cosine-versus-Euclidean silhouette comparison across K = 2 to 6",
        rows_per_chunk=12,
        font_size=8,
    )

    build_docx(figure_paths, table_paths)
    print(f"Created: {OUT_DOCX}")


if __name__ == "__main__":
    main()
