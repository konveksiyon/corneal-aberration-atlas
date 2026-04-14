import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SUPP_DIR = ROOT / "supplementary_materials"
FIG_DIR = SUPP_DIR / "figures"
TABLE_DIR = ROOT / "output" / "tables"

OUTPUT_PATH = SUPP_DIR / "BMC_Ophthalmology_Supplementary_Material_FINAL_round2.docx"

FONT_NAME = "Times New Roman"
HEADER_FILL = "D9E7F5"
ALT_FILL = "F7F9FC"
GRID_COLOR = "8A8A8A"


FIGURES = [
    (
        FIG_DIR / "Supplementary_Figure_S1.png",
        "Supplementary Figure S1. Cumulative explained variance of PCA fitted on the discovery set for the full dual-surface feature set (left) and the HOA-only feature set (right). Vertical dashed lines indicate the retained 20-component threshold; horizontal dashed lines indicate the 90% and 95% reference levels.",
        16.2,
    ),
    (
        FIG_DIR / "Supplementary_Figure_S2.png",
        "Supplementary Figure S2. Cluster-number selection metrics across K = 2 to 6 in the discovery set. Panels show cosine silhouette, bootstrap stability, severity predictability from total RMS alone, maximum pairwise centroid similarity, and the composite score used for cluster-number selection.",
        16.2,
    ),
    (
        FIG_DIR / "Supplementary_Figure_S3.png",
        "Supplementary Figure S3. Full-feature K = 4 over-segmentation atlas. The four-group probe was used to test whether forcing additional resolution produced clinically distinct morphologies; the associated pairwise post-hoc comparisons are summarized in Supplementary Table S3.",
        16.2,
    ),
    (
        FIG_DIR / "Supplementary_Figure_S4.png",
        "Supplementary Figure S4. HOA-only K = 4 over-segmentation atlas. The same probe was repeated using orders 3 to 6 only, providing an independent sensitivity analysis of the four-group solution.",
        16.2,
    ),
]


def set_run_font(run, size_pt=None, bold=None, italic=None, color=None):
    run.font.name = FONT_NAME
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), FONT_NAME)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(11)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal.element.rPr.rFonts.set(qn(f"w:{key}"), FONT_NAME)

    for style_name, size_pt in [("Heading 1", 13), ("Heading 2", 11)]:
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = Pt(size_pt)
        style.font.bold = True
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            style.element.rPr.rFonts.set(qn(f"w:{key}"), FONT_NAME)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID_COLOR, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_row_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if not any(child.tag == qn("w:tblHeader") for child in tr_pr):
        tr_pr.append(OxmlElement("w:tblHeader"))


def set_row_no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if not any(child.tag == qn("w:cantSplit") for child in tr_pr):
        tr_pr.append(OxmlElement("w:cantSplit"))


def format_caption(paragraph, caption_text):
    match = re.match(r"^(Supplementary (?:Figure|Table) S\d+\.)\s*(.*)$", caption_text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.keep_together = True
    pf.keep_with_next = False
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(3)
    pf.space_after = Pt(8)
    if match:
        run1 = paragraph.add_run(match.group(1))
        set_run_font(run1, size_pt=10, bold=True)
        run2 = paragraph.add_run(" " + match.group(2))
        set_run_font(run2, size_pt=10)
    else:
        run = paragraph.add_run(caption_text)
        set_run_font(run, size_pt=10)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.keep_with_next = True
    pf.keep_together = True
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, size_pt=16, bold=True)


def add_heading(doc, text, level=1, page_break_before=False):
    if page_break_before:
        doc.add_page_break()
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.keep_with_next = True
    pf.keep_together = True
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(10 if level == 1 else 6)
    pf.space_after = Pt(6 if level == 1 else 3)
    run = p.add_run(text)
    set_run_font(run, size_pt=13 if level == 1 else 11, bold=True)
    return p


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size_pt=11)
    return p


def add_figure(doc, image_path, caption, width_cm):
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.keep_with_next = True
    pic_p.paragraph_format.keep_together = True
    pic_p.paragraph_format.space_before = Pt(6)
    pic_p.paragraph_format.space_after = Pt(3)
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    cap = doc.add_paragraph()
    format_caption(cap, caption)


def clean_sizes(value):
    return re.sub(r"np\.int64\((\d+)\)", r"\1", str(value))


def english_interpretation(value):
    mapping = {
        "Strong separation": "Strong separation",
        "Artificial split": "Artificial split",
        "Weak separation": "Weak separation",
        "Moderate separation": "Moderate separation",
    }
    return mapping.get(value, value)


def clean_variable_name(value):
    mapping = {
        "BAD D": "BAD-D",
        "K Max (Front)": "Kmax",
    }
    return mapping.get(value, value)


def format_number(value, decimals=4, strip=True):
    if value == "—":
        return value
    if isinstance(value, str):
        return value
    text = f"{float(value):.{decimals}f}"
    if strip:
        text = text.rstrip("0").rstrip(".")
    return text


def make_table(doc, df, font_size=9.5, first_col_merge_groups=None):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    header_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        header_cells[j].text = str(col)
        shade_cell(header_cells[j], HEADER_FILL)
        set_cell_border(header_cells[j])
        header_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = header_cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        for run in p.runs:
            set_run_font(run, size_pt=font_size, bold=True)
    set_row_header(table.rows[0])
    set_row_no_split(table.rows[0])

    for i, row_vals in enumerate(df.itertuples(index=False), start=1):
        row = table.add_row()
        set_row_no_split(row)
        for j, value in enumerate(row_vals):
            cell = row.cells[j]
            cell.text = "" if pd.isna(value) else str(value)
            if i % 2 == 0:
                shade_cell(cell, ALT_FILL)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, size_pt=font_size)

    if first_col_merge_groups:
        start_row = 1
        while start_row < len(table.rows):
            end_row = min(start_row + first_col_merge_groups - 1, len(table.rows) - 1)
            merged_value = table.cell(start_row, 0).text
            merged = table.cell(start_row, 0).merge(table.cell(end_row, 0))
            merged.text = merged_value
            merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in merged.paragraphs[0].runs:
                set_run_font(run, size_pt=font_size)
            start_row += first_col_merge_groups

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    return table


def add_table_caption(doc, caption):
    p = doc.add_paragraph()
    format_caption(p, caption)
    p.paragraph_format.keep_with_next = True
    return p


def prepare_s1():
    df = pd.read_csv(TABLE_DIR / "k_comparison_extended.csv")
    out = pd.DataFrame(
        {
            "K": df["K"],
            "Cluster sizes": df["sizes"].map(clean_sizes),
            "Smallest cluster (%)": df["min_cluster_pct"].map(lambda x: format_number(x, 1)),
            "Cosine silhouette": df["cosine_silhouette"].map(lambda x: format_number(x, 4)),
            "Euclidean silhouette": df["euclidean_silhouette"].map(lambda x: format_number(x, 4)),
            "Bootstrap ARI mean": df["bootstrap_ARI_mean"].map(lambda x: format_number(x, 4)),
            "Bootstrap ARI SD": df["bootstrap_ARI_sd"].map(lambda x: format_number(x, 4)),
            "RMS-only LR accuracy": df["severity_LR_acc"].map(lambda x: format_number(x, 4)),
            "Max centroid similarity": df["max_centroid_sim"].map(lambda x: "—" if float(x) == -1.0 else format_number(x, 4)),
            "Mean centroid similarity": df["mean_centroid_sim"].map(lambda x: "—" if float(x) == -1.0 else format_number(x, 4)),
            "Chance level": df["chance_level"].map(lambda x: format_number(x, 4)),
        }
    )
    return out


def prepare_summary_table(csv_name):
    df = pd.read_csv(TABLE_DIR / csv_name)
    df["Variable"] = df["Variable"].map(clean_variable_name)
    df["Phenotype"] = df["Phenotype"].map(lambda x: f"P{x}")
    df["Median"] = df["Median"].map(lambda x: format_number(x, 2))
    df["Mean"] = df["Mean"].map(lambda x: format_number(x, 2))
    df["SD"] = df["SD"].map(lambda x: format_number(x, 2))
    return df[["Variable", "Phenotype", "N", "Median", "IQR", "Mean", "SD"]]


def prepare_pairwise_table(csv_name):
    df = pd.read_csv(TABLE_DIR / csv_name)
    df["nonsig_vars"] = df["nonsig_vars"].replace("none", "—").map(
        lambda x: x.replace("BAD D", "BAD-D").replace("K Max (Front)", "Kmax")
    )
    df["mean_cohens_d"] = df["mean_cohens_d"].map(lambda x: format_number(x, 3))
    df["max_cohens_d"] = df["max_cohens_d"].map(lambda x: format_number(x, 3))
    df["interpretation"] = df["interpretation"].map(english_interpretation)
    return df.rename(
        columns={
            "pair": "Pair",
            "n_sig": "Significant variables (n)",
            "n_nonsig": "Non-significant variables (n)",
            "nonsig_vars": "Non-significant variables",
            "mean_cohens_d": "Mean |Cohen's d|",
            "max_cohens_d": "Max |Cohen's d|",
            "interpretation": "Interpretation",
        }
    )


def prepare_s6():
    data = json.loads((TABLE_DIR / "centroid_concordance.json").read_text(encoding="utf-8"))
    matrix = pd.DataFrame(
        data["centroid_similarity_matrix"],
        columns=[f"R{idx}" for idx in range(4)],
        index=[f"D{idx}" for idx in range(4)],
    )
    matrix = matrix.map(lambda x: format_number(x, 4))
    matrix.insert(0, "Discovery", matrix.index)
    matrix = matrix.reset_index(drop=True)

    matched = pd.DataFrame(
        {
            "Discovery phenotype": [f"D{disc}" for disc, _ in data["matched_pairs_disc_to_repl"]],
            "Matched replication phenotype": [f"R{repl}" for _, repl in data["matched_pairs_disc_to_repl"]],
            "Matched cosine similarity": [format_number(x, 4) for x in data["matched_similarities"]],
        }
    )
    summary_text = (
        f"Mean matched cosine concordance = {format_number(data['mean_concordance'], 3)}; "
        f"ARI between projected and independently refit replication labels = {format_number(data['ARI_projected_vs_independent'], 2)}."
    )
    return matrix, matched, summary_text


def prepare_s7():
    df = pd.read_csv(TABLE_DIR / "weight_sensitivity.csv")
    rename = {
        "weight_scheme": "Weighting scheme",
        "w_silhouette": "Silhouette weight",
        "w_stability": "Stability weight",
        "w_severity": "Severity weight",
        "w_redundancy": "Redundancy weight",
        "composite_K2": "Composite K=2",
        "composite_K3": "Composite K=3",
        "composite_K4": "Composite K=4",
        "composite_K5": "Composite K=5",
        "composite_K6": "Composite K=6",
        "optimal_K": "Optimal K",
    }
    out = df.rename(columns=rename)
    for col in out.columns:
        if col == "Weighting scheme":
            continue
        out[col] = out[col].map(lambda x: format_number(x, 4))
    return out


def prepare_s8():
    df = pd.read_csv(TABLE_DIR / "silhouette_comparison.csv")
    out = df.rename(
        columns={
            "K": "K",
            "cosine_silhouette": "Cosine silhouette",
            "euclidean_silhouette": "Euclidean silhouette",
            "difference": "Difference",
        }
    )
    for col in out.columns[1:]:
        out[col] = out[col].map(lambda x: format_number(x, 4))
    return out


def chunk_dataframe(df, size):
    return [df.iloc[i : i + size].copy() for i in range(0, len(df), size)]


def build():
    doc = Document()
    style_document(doc)

    add_title(doc, "Supplementary Material")

    add_heading(doc, "Supplementary Methods", level=1)
    add_body_paragraph(
        doc,
        "Additional sensitivity analyses were performed using the same filtered cohort, dual-surface normalization framework, and discovery-replication split described in the main manuscript. The supplementary material reports the prespecified spherical k-means cluster-number metrics across K = 2 to 6 (Table S1; Figure S2), the full-feature and HOA-only K = 4 over-segmentation probe with post-hoc clinical summaries and pairwise distinguishability analyses (Figures S3-S4; Tables S2-S5), discovery-replication centroid concordance after centroid matching (Table S6), composite-score robustness to alternative weighting schemes (Table S7), and cosine-versus-Euclidean silhouette behavior across K = 2 to 6 (Table S8). The HOA-only feature set was used specifically to test whether the K = 4 over-segmentation pattern persisted in an independent subset of aberration features."
    )

    add_heading(doc, "Supplementary Figures", level=1, page_break_before=True)
    for idx, (image_path, caption, width_cm) in enumerate(FIGURES):
        if idx > 0:
            doc.add_page_break()
        add_figure(doc, image_path, caption, width_cm)

    add_heading(doc, "Supplementary Tables", level=1, page_break_before=True)

    # Table S1
    add_table_caption(
        doc,
        "Supplementary Table S1. Cluster-number comparison across K = 2 to 6 for the primary spherical k-means analysis in the discovery set.",
    )
    make_table(doc, prepare_s1(), font_size=8)

    # Table S2
    s2_caption = "Supplementary Table S2. Post-hoc clinical summary of the full-feature K = 4 over-segmentation probe."
    for idx, chunk in enumerate(chunk_dataframe(prepare_summary_table("clinical_summary_full.csv"), 12)):
        doc.add_page_break()
        caption = s2_caption if idx == 0 else s2_caption + " (continued)"
        add_table_caption(doc, caption)
        make_table(doc, chunk, font_size=9.5, first_col_merge_groups=4)

    # Table S3
    doc.add_page_break()
    add_table_caption(
        doc,
        "Supplementary Table S3. Pairwise distinguishability analysis for the full-feature K = 4 solution.",
    )
    make_table(doc, prepare_pairwise_table("pairwise_full.csv"), font_size=9)

    # Table S4
    s4_caption = "Supplementary Table S4. Post-hoc clinical summary of the HOA-only K = 4 over-segmentation probe."
    for idx, chunk in enumerate(chunk_dataframe(prepare_summary_table("clinical_summary_hoa.csv"), 12)):
        doc.add_page_break()
        caption = s4_caption if idx == 0 else s4_caption + " (continued)"
        add_table_caption(doc, caption)
        make_table(doc, chunk, font_size=9.5, first_col_merge_groups=4)

    # Table S5
    doc.add_page_break()
    add_table_caption(
        doc,
        "Supplementary Table S5. Pairwise distinguishability analysis for the HOA-only K = 4 solution.",
    )
    make_table(doc, prepare_pairwise_table("pairwise_hoa.csv"), font_size=9)

    # Table S6
    doc.add_page_break()
    add_table_caption(
        doc,
        "Supplementary Table S6. Discovery-replication centroid concordance after centroid matching for the K = 4 over-segmentation probe.",
    )
    matrix_df, matched_df, summary_text = prepare_s6()
    note = doc.add_paragraph("Cosine similarity matrix between discovery (D rows) and replication (R columns) centroids.")
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    note.paragraph_format.space_after = Pt(4)
    set_run_font(note.runs[0], size_pt=9.5, italic=True)
    make_table(doc, matrix_df, font_size=9)

    note2 = doc.add_paragraph("Matched centroid pairs after one-to-one assignment.")
    note2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    note2.paragraph_format.space_before = Pt(6)
    note2.paragraph_format.space_after = Pt(4)
    set_run_font(note2.runs[0], size_pt=9.5, italic=True)
    make_table(doc, matched_df, font_size=9.5)
    note3 = doc.add_paragraph(summary_text)
    note3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    note3.paragraph_format.space_before = Pt(4)
    note3.paragraph_format.space_after = Pt(0)
    set_run_font(note3.runs[0], size_pt=9.5)

    # Table S7
    doc.add_page_break()
    add_table_caption(
        doc,
        "Supplementary Table S7. Composite-score sensitivity to alternative weighting schemes across K = 2 to 6.",
    )
    make_table(doc, prepare_s7(), font_size=8.5)

    # Table S8
    doc.add_page_break()
    add_table_caption(
        doc,
        "Supplementary Table S8. Cosine-versus-Euclidean silhouette comparison across K = 2 to 6.",
    )
    make_table(doc, prepare_s8(), font_size=9.5)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
